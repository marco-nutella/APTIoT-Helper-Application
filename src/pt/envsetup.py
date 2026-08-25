import flet as ft
import base64
import io
from docx import Document


class EnvSetupView:
    def __init__(self, page):
        self.page = page
        self.document = Document()
        
        self.file_picker = ft.FilePicker()
        self.image_src:bytes
        self.image = ft.Image(
                src="",
                expand=True,
                width=600,
                height=600,
                border_radius=5,
                visible=False,
            )
        self.next_button = ft.Button("Next Step: Intelligence Gathering", on_click=self.next_step, style=ft.ButtonStyle(
                        color={
                            ft.ControlState.HOVERED: ft.Colors.WHITE,
                            ft.ControlState.DEFAULT: ft.Colors.BLACK,
                        },
                        bgcolor={
                            ft.ControlState.HOVERED: ft.Colors.GREEN_100,
                            ft.ControlState.DEFAULT: ft.Colors.GREEN_400,
                        },
                    ), visible=False)

    def write_document(self, pt_handler):
        self.document = pt_handler.create_doc_header(self.document)
        
        self.document.add_heading("Environment Setup & Familiarization", level=1)
        self.document.add_heading("Local Network Diagram of Environment Setup", level=2)

        diagram_bytes_stream = io.BytesIO(self.image_src)
        section = self.document.sections[0]
        self.document.add_paragraph().add_run().add_picture(diagram_bytes_stream, width=section.page_width-section.left_margin-section.right_margin) # type: ignore
        self.document.save(f"{pt_handler.device.name} Environment Setup.docx")
        pt_handler.documents["envsetup"] = self.document

    def next_step(self):
        pt_handler = self.page.session.store.get("pt_handler")
        self.write_document(pt_handler)
        self.page.session.store.set("pt_handler", pt_handler)
        pt_handler.open_intelgathering()

    async def upload_images(self):
        images = await self.file_picker.pick_files(file_type=ft.FilePickerFileType.IMAGE, allow_multiple=False, with_data=True)
        for i in images:
            if not i.bytes:
                continue

            self.image_src = i.bytes
            self.image.src = base64.b64encode(i.bytes).decode("utf-8")
            self.image.visible = True
            self.next_button.visible = True
            self.page.update()

    def populate(self):
            return ft.Container(
                alignment = ft.Alignment.CENTER,
                content= ft.Column(
                    margin = 15,
                    spacing = 15,
                    controls=[
                        ft.Row(
                            controls=[
                            ft.Text("Environment Setup & Familiarization", size=36, weight=ft.FontWeight.W_800),
                            self.next_button,
                            ]
                        ),
                        ft.Column(
                            scroll = ft.ScrollMode.AUTO,
                            alignment=ft.MainAxisAlignment.START,
                            expand = True,
                            controls=[
                                ft.Text(value="Description", expand=True, size=24, weight=ft.FontWeight.W_600),
                                ft.Text(value="""Familiarize yourself with the device to be tested and its functions. Read its manual, install any applications it might interface with,
and configure it in your testing environment. Your testing environment should accomodate for as many functionalities of the device as is
feasible for you and/or your organization. As a rule of thumb, if the device is interoperable with others and it is possible to have them 
be present in your testing environment, then they should be. Make a local network diagram of your testing environment and submit it to proceed.
"""
                                , expand=True, size=16, weight=ft.FontWeight.W_400, text_align=ft.TextAlign.LEFT), # """ Breaks spacing in the string, so these lines needs to break identation.
                                ft.Text(value="Deliverables", expand=True, size=36, weight=ft.FontWeight.W_800),
                                ft.Text(value="Local Network Diagram of Environment Setup", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.image,
                                ft.Button("Add Diagram Image", on_click=self.upload_images),
                            ]
                        )
                    ]
                )
            )