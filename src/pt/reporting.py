import flet as ft
import base64
import io
import os
from docx import Document
from idses.idsesutilities import Protocols
from idses.idsescalculator import IDSESWidget

class ReportingView:
    def __init__(self, page): # We define self.device later, since it's not yet defined by the time the object gets instantiated.
        self.page = page
        self.document = Document()

        self.summary_input = ft.TextField(label="Penetration Test Summary", hint_text="Enter a summary of exploited vulnerabilities and general findings here...", width=800, multiline=True, min_lines=1, max_lines=20, on_change=self.make_next_visible)
        self.security_recommendations_input = ft.TextField(label="Recommendations for Device Security", hint_text="Enter recommendations for improvement of device security here...", width=800, multiline=True, min_lines=1, max_lines=20)
        self.posture_recommendation_input = ft.TextField(label="Recommendations for Security Posture", hint_text="Enter recommendations for improvement of general security posture here...", width=800, multiline=True, min_lines=1, max_lines=20)
        
        self.next_button = ft.Button("Finish", on_click=self.next_step, style=ft.ButtonStyle(
                        color={
                            ft.ControlState.HOVERED: ft.Colors.WHITE,
                            ft.ControlState.DEFAULT: ft.Colors.BLACK,
                        },
                        bgcolor={
                            ft.ControlState.HOVERED: ft.Colors.GREEN_100,
                            ft.ControlState.DEFAULT: ft.Colors.GREEN_400,
                        },
                    ), visible=False)

    def make_next_visible(self):
        self.next_button.visible = True

    def write_document(self, pt_handler):
        self.document = pt_handler.create_doc_header(self.document)

        device = self.page.session.store.get("pt_handler").device
        self.document.add_heading("Device Final Absolute IDSES Score:      ", level=0).add_run(f"{device.get_idses_score()}, Weighted Score: {device.get_weighted_idses_score()}")

        self.document.add_heading("Reporting", level=1)
        self.document.add_heading("Penetration Test Summary", level=2)
        self.document.add_paragraph(self.summary_input.value)

        self.document.add_heading("Recommendations for Device Security", level=2)
        self.document.add_paragraph(self.security_recommendations_input.value)

        self.document.add_heading("Recommendations for Security Posture", level=2)
        self.document.add_paragraph(self.posture_recommendation_input.value)

        self.document.save(os.path.join(self.page.session.store.get("pt_handler").save_path, f"{pt_handler.device.name} Reporting.docx"))
        pt_handler.documents["reporting"] = self.document

    def next_step(self):
        pt_handler = self.page.session.store.get("pt_handler")
        self.write_document(pt_handler)
        self.page.session.store.set("pt_handler", pt_handler)
        pt_handler.open_final()

    def populate(self):
            self.idses_widget = IDSESWidget(self.page.session.store.get("pt_handler").device)
            return ft.Container(
                alignment = ft.Alignment.CENTER,
                content= ft.Column(
                    margin = 15,
                    spacing = 15,
                    controls=[
                        ft.Row(
                            controls=[
                            ft.Text("Reporting", size=36, weight=ft.FontWeight.W_800),
                            self.next_button,
                            ]
                        ),
                        ft.Column(
                            scroll = ft.ScrollMode.AUTO,
                            alignment=ft.MainAxisAlignment.START,
                            expand = True,
                            controls=[
                                ft.Text(value="Description", expand=True, size=24, weight=ft.FontWeight.W_600),
                                ft.Text(value="""Conduct Open Source Intelligence (OSINT) on the device you are testing, its vendor/manufacturer, and previous security incidents involving either of them.
Map out the protocols and technologies the device makes use of. The device's manual is likely to already have some of this information. Research other similar products by the same vendor, and try to understand
if their devices share technologies, or are otherwise part of an ecosystem. Finally, search for previous security incidents involving either this device or others by the same vendor. The CVE database at cve.org
might contain information on previous vulnerabilities."""
                                , expand=True, size=16, weight=ft.FontWeight.W_400, text_align=ft.TextAlign.LEFT), # """ Breaks spacing in the string, so these lines needs to break identation.
                                ft.Text(value="Final IDSES Score", expand=True, size=36, weight=ft.FontWeight.W_800),
                                self.idses_widget.populate(),
                                ft.Text(value="Deliverables", expand=True, size=36, weight=ft.FontWeight.W_800),
                                ft.Text(value="Penetration Test Summary", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.summary_input,
                                ft.Text(value="Recommendations for Device Security", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.security_recommendations_input,
                                ft.Text(value="Recommendations for Security Posture", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.posture_recommendation_input,
                            ]
                        )
                    ]
                )
            )