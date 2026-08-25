import flet as ft
from idses.idsescalculator import IDSESDevice
from idses.idsesutilities import RiskEnvironment, RiskCriticality
from docx import Document
import io

class Organization:
    def __init__(self, name:str, logo:bytes|None = None):
        self.name = name
        self.logo = logo

class APTIoTPenetrationTest:
    def __init__(self, page, device:IDSESDevice, org:Organization, testers:str, date:str):
        self.page = page
        self.device = device
        self.org = org
        self.testers = testers
        self.date = date
        self.documents = {
            "envsetup": None,
            "intelgathering": None,
            "trafficanalysis": None,
            "vulnassessment": None,
            "exploitation": None,
            "reporting": None,
            "finalreport": None,
        }

    def get_org_info(self) -> tuple[str, bytes|None]:
        return self.org.name, self.org.logo

    def get_pt_string(self) -> str:
        return f"Penetration Test of IoT Device '{self.device.name}' ({self.device.vendor}, {self.device.year}), by tester(s) {self.testers} on {self.date}."

    def quit(self):
        self.page.navigate("")

    def open_envsetup(self):
        self.page.navigate("/envsetup")

    def open_intelgathering(self):
        self.page.navigate("/intelgathering")

    def open_trafficanalysis(self):
        self.page.navigate("/trafficanalysis")

    def open_vulnassessment(self):
        self.page.navigate("/vulnassessment")

    def open_exploitation(self):
        self.page.navigate("/exploitation")

    def open_reporting(self):
        self.page.navigate("/reporting")

    def create_doc_header(self, document):
        org = self.get_org_info()
        top_heading = document.add_heading(f"{org[0]}", 0)
        document.add_heading(self.get_pt_string(), level=2)

        if org[1]:
            image_bytes_stream = io.BytesIO(org[1])
            section = document.sections[0]
            top_heading.add_run().add_picture(image_bytes_stream, width=(section.page_width-section.left_margin-section.right_margin)/2) # type: ignore # Measurements in English Metric Units (Emus). We don't want the image to be too big and not fit.

        return document

class PTAssistantMainTab:
    def __init__(self, page):
        self.page = page
        self.device = IDSESDevice()

        self.org_name_input = ft.TextField(label="Organization Name", hint_text="Enter the organization's name here...", width=400)
        self.tester_name_input = ft.TextField(label="Tester Name(s)", hint_text="Enter your name(s) here...", width=400)
        self.date_input = ft.TextField(label="Enter Date", hint_text="Suggested format: MM/YYYY", width=400)

        self.name_input = ft.TextField(label="Device Name", hint_text="Enter the device's name here...", width=400, on_change=self.on_change_device_name)
        self.vendor_input = ft.TextField(label="Vendor Name", hint_text="Enter the device's vendor here...", width=400, multiline=True, min_lines=1, max_lines=10, on_change=self.on_change_device_vendor)
        self.year_input = ft.TextField(label="Enter Year", hint_text="...", width=100, on_change=self.on_change_device_year,
                        input_filter=ft.InputFilter(
                            regex_string=r"^\d+$",
                            allow=True,
                            replacement_string=""
                        )
                    )

        self.environment_dropdown = ft.Dropdown(
                        width = 300,
                        key = "environment_dropdown",
                        label = "Select...",
                        options = [
                            ft.DropdownOption(
                                key = env.value[0],
                                content = ft.Text(value=env.value[0], size=14, weight=ft.FontWeight.W_400),
                            )
                            for env in RiskEnvironment
                        ],
                        on_select = self.on_select_environment
                    )

        self.criticality_dropdown = ft.Dropdown(
                                width = 300,
                                key = "environment_dropdown",
                                label = "Select...",
                                options = [
                                    ft.DropdownOption(
                                        key = env.value[0],
                                        content = ft.Text(value=env.value[0], size=14, weight=ft.FontWeight.W_400),
                                    )
                                    for env in RiskCriticality
                                ],
                                on_select = self.on_select_criticality
                            )
        
        self.file_picker = ft.FilePicker()
        self.checkmark = ft.Icon(ft.Icons.CHECK, color=ft.Colors.GREEN_400, size=24, visible=False)
        self.org_logo: bytes|None = None

    def on_change_device_name(self, event:ft.Event[ft.TextField]):
        self.device.update_information(name=event.control.value)
    
    def on_change_device_vendor(self, event:ft.Event[ft.TextField]):
        self.device.update_information(manufacturer=event.control.value)

    def on_change_device_year(self, event:ft.Event[ft.TextField]):
        self.device.update_information(year=int(event.control.value))

    def on_select_environment(self, event:ft.Event[ft.Dropdown]):
        value = str(event.control.value).upper()
        self.device.update_risk_environment(RiskEnvironment[value])

    def on_select_criticality(self, event:ft.Event[ft.Dropdown]):
        value = str(event.control.value).upper()
        self.device.update_risk_criticality(RiskCriticality[value])

    async def upload_images(self):
        images = await self.file_picker.pick_files(file_type=ft.FilePickerFileType.IMAGE, allow_multiple=False, with_data=True)
        for i in images:
            if not i.bytes:
                continue

            self.org_logo = i.bytes
            self.checkmark.visible = True

    def start_penetration_test(self):
        if not self.org_name_input.value or not self.tester_name_input.value or not self.date_input.value:
            self.page.show_dialog(ft.SnackBar("Incomplete penetration testing information. Please verify and try again."))
            return
        org = Organization(self.org_name_input.value, self.org_logo)

        if not self.name_input.value or not self.vendor_input.value or not self.year_input.value or not self.environment_dropdown.value or not self.criticality_dropdown.value:
            self.page.show_dialog(ft.SnackBar("Incomplete device information. Please verify and try again."))
            return

        pt_handler = APTIoTPenetrationTest(self.page, self.device, org, self.tester_name_input.value, self.date_input.value)
        self.page.session.store.set("pt_handler", pt_handler)

        pt_handler.open_envsetup()


    def populate(self):
        return ft.Container(
            alignment = ft.Alignment.CENTER,
            expand=True,
            content= ft.Column(
                margin = 15,
                spacing = 15,
                scroll = ft.ScrollMode.AUTO,
                alignment=ft.MainAxisAlignment.START,
                controls=[
                    ft.Text("Penetration Test Information", size=36, weight=ft.FontWeight.W_800),
                    ft.Text(value="Organization Name", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.org_name_input,
                    ft.Text(value="Tester Name(s)", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.tester_name_input,
                    ft.Text(value="Testing Date", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.date_input,
                    ft.Text(value="(Optional) Organization Logo", expand=True, size=24, weight=ft.FontWeight.W_600),
                    ft.Row(
                        controls=[
                            ft.Button("Add Logo", on_click=self.upload_images),
                            self.checkmark
                        ]
                    ),
                    ft.Text("Device Information", size=36, weight=ft.FontWeight.W_800),
                    ft.Text(value="Name", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.name_input,
                    ft.Text(value="Vendor", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.vendor_input,
                    ft.Text(value="Release Year", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.year_input,
                    ft.Text("Risk Metrics", size=36, weight=ft.FontWeight.W_800),
                    ft.Text(value="Environment", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.environment_dropdown,
                    ft.Text(value="Criticality", expand=True, size=24, weight=ft.FontWeight.W_600),
                    self.criticality_dropdown,
                    ft.Button("Start Penetration Test", on_click=self.start_penetration_test)
                ]
            )
        )
    