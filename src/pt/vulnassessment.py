import flet as ft
import base64
import io
import os
from functools import partial
from docx import Document
from ivss.ivssvulnerability import IVSSVulnerability

class IVSSStubWidget:
    def __init__(self, page, vuln:IVSSVulnerability, edit_mode:bool = True):
        self.page = page
        self.ivss_vulnerability = vuln
        self.edit_mode = edit_mode

        self.file_picker = ft.FilePicker()
        self.images_src = []
        self.images_display = ft.Row(
                        expand=True,
                        scroll=ft.ScrollMode.AUTO,
                        width=800,
                        controls=[]
                    )
        self.display_images_display = ft.Row(
                        expand=True,
                        scroll=ft.ScrollMode.AUTO,
                        width=800,
                        controls=[]
                    )

        self.name_input = ft.TextField(label="Vulnerability Name", hint_text="Enter a vulnerability name here...", width=400, on_change=self.on_change_vulnerability_name)
        self.protocol_input = ft.TextField(label="Vulnerability Protocol", hint_text="Enter the vulnerable protocol here...", width=200, on_change=self.on_change_vulnerability_protocol)
        self.description_input = ft.TextField(label="Vulnerability Description", hint_text="Enter vulnerability details here...", width=800, multiline=True, min_lines=1, max_lines=10, on_change=self.on_change_vulnerability_description)
        self.year_input = ft.TextField(label="Enter Year", hint_text="...", width=100, on_change=self.on_change_vulnerability_year,
                        input_filter=ft.InputFilter(
                            regex_string=r"^\d+$",
                            allow=True,
                            replacement_string=""
                        )
                    )

        self.protocol_display = ft.Text(self.ivss_vulnerability.get_protocol(), expand=True, size=36, weight=ft.FontWeight.W_800)
        self.name_display = ft.Text(self.ivss_vulnerability.get_name(), expand=True, size=24, weight=ft.FontWeight.W_700)
        self.year_display = ft.Text(str(self.ivss_vulnerability.get_year()), expand=True, size=20, weight=ft.FontWeight.W_700)
        self.description_display = ft.Text(self.ivss_vulnerability.get_description(), expand=True, size=16, weight=ft.FontWeight.W_600)

        self.edit_widget = ft.Column(
            controls=[
                ft.Row(
                    controls=[
                        self.name_input,
                        self.protocol_input,
                        self.year_input
                    ]
                ),
                self.description_input,
                self.images_display,
                ft.Button("Add Images", on_click=self.upload_images),
                ft.Button("Save Changes", on_click=self.change_mode, style=ft.ButtonStyle(
                        color={
                            ft.ControlState.HOVERED: ft.Colors.WHITE,
                            ft.ControlState.DEFAULT: ft.Colors.BLACK,
                        },
                        bgcolor={
                            ft.ControlState.HOVERED: ft.Colors.GREEN_100,
                            ft.ControlState.DEFAULT: ft.Colors.GREEN_400,
                        },
                    )),
            ],
            visible = self.edit_mode
        )

        self.display_widget = ft.Column(
            controls=[
                ft.Row(
                    alignment=ft.MainAxisAlignment.START,
                    controls=[
                        self.protocol_display,
                        self.name_display,
                        self.year_display
                    ]
                ),
                self.description_display,
                self.display_images_display,
                ft.Button("Edit Vulnerability", on_click=self.change_mode),
            ],
            visible = not self.edit_mode
        )
        
    def update_vulnerability_details(self, name:str = "", description:str = "", year:int = 0, images:list[ft.FilePickerFile] = []):
        self.ivss_vulnerability.update_info(name, description, year)
        if images:
            for i in images:
                if not i.bytes:
                    continue
                i.path = None
            self.ivss_vulnerability.add_images(images)

    def on_change_vulnerability_name(self, event:ft.Event[ft.TextField]):
        self.update_vulnerability_details(name=event.control.value)
        self.name_display.value = self.ivss_vulnerability.get_name()

    def on_change_vulnerability_description(self, event:ft.Event[ft.TextField]):
        self.update_vulnerability_details(description=event.control.value)
        self.description_display.value = self.ivss_vulnerability.get_description()

    def on_change_vulnerability_year(self, event:ft.Event[ft.TextField]):
        self.update_vulnerability_details(year=int(event.control.value))
        self.year_display.value = str(self.ivss_vulnerability.get_year())

    def on_change_vulnerability_protocol(self, event:ft.Event[ft.TextField]):
        self.ivss_vulnerability.set_protocol(event.control.value)
        self.protocol_display.value = self.ivss_vulnerability.get_protocol()

    def change_mode(self):
        self.edit_mode = not self.edit_mode
        self.edit_widget.visible = self.edit_mode
        self.display_widget.visible = not self.edit_mode
        self.page.update()

    def get_images_src(self) -> list[bytes]:
        return self.images_src

    def get_vulnerability(self) -> IVSSVulnerability:
        return self.ivss_vulnerability

    async def upload_images(self):
        images = await self.file_picker.pick_files(file_type=ft.FilePickerFileType.CUSTOM, allowed_extensions=["png", "jpg", "jpeg"], allow_multiple=True, with_data=True)
        for i in images:
            if not i.bytes:
                continue

            self.images_src.append(i.bytes)
            for v in [self.images_display, self.display_images_display]:
                v.controls.append(
                    ft.IconButton(
                        icon=(
                            ft.Image(
                                src=base64.b64encode(i.bytes).decode("utf-8"),
                                expand=True,
                                width=200,
                                height=200,
                                border_radius=5,
                            )
                        ),
                        on_click=partial(self.fullscreen_image, base64.b64encode(i.bytes).decode("utf-8"), self.page),
                ))

        self.page.update()
        self.update_vulnerability_details(images=images)

    def fullscreen_image(self, image, page):
        img_container = ft.Container(
            content=ft.Image(
                src=image,
            ),
            expand=True,
        )
        
        dialog = ft.AlertDialog(
            content=img_container,
            actions=[
                ft.TextButton("Close", on_click=lambda _: page.pop_dialog()),
            ],
            modal=True,
        )
        
        page.show_dialog(dialog)

    def populate(self) -> ft.Container:
        return ft.Container(
            expand=True,
            border_radius=50,
            shadow=ft.BoxShadow(
                spread_radius=1,
                blur_radius=3,
                color=ft.Colors.BLUE_GREY_300,
                offset=ft.Offset(0, 0),
                blur_style=ft.BlurStyle.OUTER,
            ),
            content= ft.Column(
                margin = 15,
                spacing = 15,
                controls=[
                    self.edit_widget,
                    self.display_widget,
                ]
            )
        )

class VulnAssessmentView:
    def __init__(self, page):
        self.page = page
        self.document = Document()
        self.vulnerability_stubs = []

        self.ports_input = ft.TextField(label="Open Ports and Available Services", hint_text="Enter the open ports and available services in this device...", width=800, multiline=True, min_lines=1, max_lines=20)
        self.tools_input = ft.TextField(label="Tools and Software Used", hint_text="Enter the tools and software that were used for this step...", width=400, multiline=True, min_lines=1, max_lines=4)
        
        self.next_button = ft.Button("Next Step: Exploitation", on_click=self.next_step, style=ft.ButtonStyle(
                        color={
                            ft.ControlState.HOVERED: ft.Colors.WHITE,
                            ft.ControlState.DEFAULT: ft.Colors.BLACK,
                        },
                        bgcolor={
                            ft.ControlState.HOVERED: ft.Colors.GREEN_100,
                            ft.ControlState.DEFAULT: ft.Colors.GREEN_400,
                        },
                    ), visible=False)
        
        self.add_vulnerability_button = ft.Button("Add New Vulnerability", on_click=self.add_new_vulnerability)
        self.vulnerability_stubs_widgets = []

    def write_document(self, pt_handler):
        self.document = pt_handler.create_doc_header(self.document)

        self.document.add_heading("Vulnerability Assessment", level=1)
        self.document.add_heading("Open Ports and Available Services", level=2)
        self.document.add_paragraph(self.ports_input.value)
        self.document.add_heading("Tools and Software Used", level=2)
        self.document.add_paragraph(self.tools_input.value)
        self.document.add_heading("Vulnerabilities", level=2)

        section = self.document.sections[0]
        for i,v in pt_handler.device.vulnerabilities.items(): # We get vulnerabilities from the device, as they have updated IDs there.
            self.document.add_heading(f"{i} ({v.get_year()})", level=3)
            self.document.add_paragraph().add_run(f"ID: {v.get_id()}").bold = True
            self.document.add_paragraph(v.get_description())
            self.document.add_heading("Images", level=4)

            for i in v.images:
                diagram_bytes_stream = io.BytesIO(i.bytes)
                self.document.add_paragraph().add_run().add_picture(diagram_bytes_stream, width=section.page_width-section.left_margin-section.right_margin) # type: ignore

        self.document.save(os.path.join(self.page.session.store.get("pt_handler").save_path, f"{pt_handler.device.name} Vulnerability Assessment.docx"))
        pt_handler.documents["vulnassessment"] = self.document
        self.page.session.store.set("pt_handler", pt_handler)

    def next_step(self):
        pt_handler = self.page.session.store.get("pt_handler")
        device = self.page.session.store.get("pt_handler").device
        for v in self.vulnerability_stubs:
            vuln = v.get_vulnerability()
            device.add_vulnerability(vuln)

        self.page.session.store.set("pt_handler", pt_handler)
        self.write_document(pt_handler)
        pt_handler.open_exploitation()

    def add_new_vulnerability(self):
        new_vuln = IVSSVulnerability()
        vuln_stub = IVSSStubWidget(self.page, new_vuln)
        self.vulnerability_stubs.append(vuln_stub)
        self.vulnerability_stubs_widgets.append(vuln_stub.populate())
        self.next_button.visible = True
        self.page.update()

    def populate(self):
            return ft.Container(
                alignment = ft.Alignment.CENTER,
                content= ft.Column(
                    margin = 15,
                    spacing = 15,
                    controls=[
                        ft.Row(
                            controls=[
                            ft.Text("Vulnerability Assessment", size=36, weight=ft.FontWeight.W_800),
                            self.next_button,
                            ]
                        ),
                        ft.Column(
                            scroll = ft.ScrollMode.AUTO,
                            alignment=ft.MainAxisAlignment.START,
                            expand = True,
                            controls=[
                                ft.Text(value="Description", expand=True, size=24, weight=ft.FontWeight.W_600),
                                ft.Text(value="""Vulnerability Assessment is the first of the two most critical parts of APTIoT. In it, the goal is to use software (and potentially hardware) tools
to investigate potential vulnerabilities in the device, making use of all of the information acquired up to this point. Vulnerabilities aren't scored in this step, and don't have to be confirmed - 
treat it as a list of vulnerabilities that should be verified in the Exploitation step. Document the open ports and available services in the device, as well as the tools that were used in this process,
before listing all relevant potential vulnerabilities."""
                                , expand=True, size=16, weight=ft.FontWeight.W_400, text_align=ft.TextAlign.LEFT), # """ Breaks spacing in the string, so these lines needs to break identation.
                                ft.Text(value="Deliverables", expand=True, size=36, weight=ft.FontWeight.W_800),
                                ft.Text(value="Open Ports and Available Services", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.ports_input,
                                ft.Text(value="Tools and Software Used", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.tools_input,
                                ft.Row(
                                    alignment=ft.MainAxisAlignment.START,
                                    controls=[
                                        ft.Text(value="Identified Vulnerabilities", size=24, weight=ft.FontWeight.W_600),
                                        self.add_vulnerability_button,
                                    ]
                                ),
                                ft.Column(
                                    expand=True,
                                    scroll=ft.ScrollMode.AUTO,
                                    height=800,
                                    margin=15,
                                    spacing=15,
                                    controls=self.vulnerability_stubs_widgets
                                )
                            ]
                        )
                    ]
                )
            )