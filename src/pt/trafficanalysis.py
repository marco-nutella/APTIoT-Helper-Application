import flet as ft
import base64
import io
from docx import Document

class TrafficAnalysisView:
    def __init__(self, page):
        self.page = page
        self.document = Document()
        
        self.file_picker = ft.FilePicker()
        self.images_src = []
        self.images_display = ft.Row(
                        scroll=ft.ScrollMode.AUTO,
                        width=800,
                        expand=True,
                        controls=[]
                    )
        self.notes_input = ft.TextField(label="Additional Notes", hint_text="Enter any additional remarks from this step here...", width=800, multiline=True, min_lines=1, max_lines=20)
        
        self.next_button = ft.Button("SKIP to Next Step: Vulnerability Assessment", on_click=self.next_step, style=ft.ButtonStyle(
                        color={
                            ft.ControlState.HOVERED: ft.Colors.WHITE,
                            ft.ControlState.DEFAULT: ft.Colors.BLACK,
                        },
                        bgcolor={
                            ft.ControlState.HOVERED: ft.Colors.RED_100,
                            ft.ControlState.DEFAULT: ft.Colors.RED_400,
                        },
                    ), visible=True)

    def write_document(self, pt_handler):
        self.document = pt_handler.create_doc_header(self.document)

        self.document.add_heading("Traffic Analysis", level=1)
        self.document.add_heading("Network Diagrams of Relevant Device Communications", level=2)

        if self.images_src:
            section = self.document.sections[0]
            for i in self.images_src:
                diagram_bytes_stream = io.BytesIO(i)
                self.document.add_paragraph().add_run().add_picture(diagram_bytes_stream, width=section.page_width-section.left_margin-section.right_margin) # type: ignore
            self.document.add_heading("Additional Notes", level=2)
            self.document.add_paragraph(self.notes_input.value)
        else:
            self.document.add_heading("Skipped by the tester(s).", level=1)

        self.document.save(f"{pt_handler.device.name} Traffic Analysis.docx")
        pt_handler.documents["trafficanalysis"] = self.document

    def next_step(self):
        pt_handler = self.page.session.store.get("pt_handler")
        self.write_document(pt_handler)
        self.page.session.store.set("pt_handler", pt_handler)
        pt_handler.open_vulnassessment()

    async def upload_images(self):
        images = await self.file_picker.pick_files(file_type=ft.FilePickerFileType.IMAGE, allow_multiple=True, with_data=True)
        for i in images:
            if not i.bytes:
                continue

            self.images_src.append(i.bytes)
            self.images_display.controls.append(ft.Image(
                src=base64.b64encode(i.bytes).decode("utf-8"),
                expand=True,
                width=600,
                height=600,
                border_radius=5,
            ))

        self.next_button.content = "Next Step: Vulnerability Assessment"
        self.next_button.style = ft.ButtonStyle(
            color={
                ft.ControlState.HOVERED: ft.Colors.WHITE,
                ft.ControlState.DEFAULT: ft.Colors.BLACK,
            },
            bgcolor={
                ft.ControlState.HOVERED: ft.Colors.GREEN_100,
                ft.ControlState.DEFAULT: ft.Colors.GREEN_400,
            },
        )
        self.page.update()

    def populate(self):
            return ft.Container(
                alignment = ft.Alignment.CENTER,
                content= ft.Column(
                    margin = 15,
                    spacing = 15,
                    controls=[
                        ft.Row(
                            controls=[
                            ft.Text("(Optional) Traffic Analysis", size=36, weight=ft.FontWeight.W_800),
                            self.next_button,
                            ]
                        ),
                        ft.Column(
                            scroll = ft.ScrollMode.AUTO,
                            alignment=ft.MainAxisAlignment.START,
                            expand = True,
                            controls=[
                                ft.Text(value="Description", expand=True, size=24, weight=ft.FontWeight.W_600),
                                ft.Text(value="""(Currently) Optional step revolving the documentation of wired or wireless communications the device is capable of performing. Traffic Analysis
consists of using some manner of network monitor (such as WireShark) in order to capture exchanges and transactions between devices for further manual analysis. In doing so, the tester is able to
understand the implementation of the device\'s features, as well as study possible vulnerabilities, such as a lack of/weak encryption in specific transactions. The result is a series of network communication diagrams
(think of a TCP 3-way handshake diagram) documenting different types of exchanges between devices."""
                                , expand=True, size=16, weight=ft.FontWeight.W_400, text_align=ft.TextAlign.LEFT), # """ Breaks spacing in the string, so these lines needs to break identation.
                                ft.Text(value="Deliverables", expand=True, size=36, weight=ft.FontWeight.W_800),
                                ft.Text(value="Network Diagrams of Relevant Device Communications", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.images_display,
                                ft.Button("Add Diagram Image(s)", on_click=self.upload_images),
                                ft.Text(value="Additional Notes", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.notes_input,
                            ]
                        )
                    ]
                )
            )