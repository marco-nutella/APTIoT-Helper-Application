import flet as ft
import base64
import io
import os
from docx import Document
from idses.idsesutilities import Protocols

class IntelGatheringView:
    def __init__(self, page): # We define self.device later, since it's not yet defined by the time the object gets instantiated.
        self.page = page
        self.document = Document()

        self.checkboxes = []
        self.device_text_input = ft.TextField(label="Device Research", hint_text="Enter device OSINT research here...", width=800, multiline=True, min_lines=1, max_lines=20)
        self.vendor_text_input = ft.TextField(label="Vendor Research", hint_text="Enter vendor OSINT research here...", width=800, multiline=True, min_lines=1, max_lines=20)
        self.security_text_input = ft.TextField(label="Security Research", hint_text="Enter security OSINT research here...", width=800, multiline=True, min_lines=1, max_lines=20)
        for p in Protocols: 
            self.checkboxes.append(ft.Checkbox(label=p.value[0], value=False, data=p, on_change=self.on_change_device_protocols))
        
        self.next_button = ft.Button("Next Step: Traffic Analysis", on_click=self.next_step, style=ft.ButtonStyle(
                        color={
                            ft.ControlState.HOVERED: ft.Colors.WHITE,
                            ft.ControlState.DEFAULT: ft.Colors.BLACK,
                        },
                        bgcolor={
                            ft.ControlState.HOVERED: ft.Colors.GREEN_100,
                            ft.ControlState.DEFAULT: ft.Colors.GREEN_400,
                        },
                    ), visible=False)

    def write_document(self, pt_handler):
        self.document = pt_handler.create_doc_header(self.document)

        self.document.add_heading("Intelligence Gathering", level=1)
        self.document.add_heading("Device Research", level=2)
        self.document.add_paragraph(self.device_text_input.value)

        self.document.add_heading("Vendor Research", level=2)
        self.document.add_paragraph(self.vendor_text_input.value)

        self.document.add_heading("Security Research", level=2)
        self.document.add_paragraph(self.security_text_input.value)

        self.document.save(os.path.join(self.page.session.store.get("pt_handler").save_path, f"{pt_handler.device.name} Intelligence Gathering.docx"))
        pt_handler.documents["intelgathering"] = self.document

    def next_step(self):
        pt_handler = self.page.session.store.get("pt_handler")
        pt_handler.device = self.device
        self.write_document(pt_handler)
        self.page.session.store.set("pt_handler", pt_handler)
        pt_handler.open_trafficanalysis()

    def on_change_device_protocols(self, event:ft.Event[ft.Checkbox]):
            self.next_button.visible = True
            if event.control.value:
                self.device.add_protocol(event.control.data)
            else:
                self.device.remove_protocol(event.control.data)

    def populate(self):
            self.device = self.page.session.store.get("pt_handler").device
            checkboxes_split_point = round(len(self.checkboxes)/2)
            return ft.Container(
                alignment = ft.Alignment.CENTER,
                content= ft.Column(
                    margin = 15,
                    spacing = 15,
                    controls=[
                        ft.Row(
                            controls=[
                            ft.Text("Intelligence Gathering", size=36, weight=ft.FontWeight.W_800),
                            self.next_button,
                            ]
                        ),
                        ft.Column(
                            scroll = ft.ScrollMode.AUTO,
                            alignment=ft.MainAxisAlignment.START,
                            expand = True,
                            controls=[
                                ft.Text(value="Description", expand=True, size=24, weight=ft.FontWeight.W_600),
                                ft.Text(value="""Conduct Open Source Intelligence (OSINT) on the device you are testing, its vendor/manufacturer, and previous security incidents involving either of them.
Map out the protocols and technologies the device makes use of. The device's manual is likely to already have some of this information. Research other similar products by the same vendor, and try to understand
if their devices share technologies, or are otherwise part of an ecosystem. Finally, search for previous security incidents involving either this device or others by the same vendor. The CVE database at cve.org
might contain information on previous vulnerabilities."""
                                , expand=True, size=16, weight=ft.FontWeight.W_400, text_align=ft.TextAlign.LEFT), # """ Breaks spacing in the string, so these lines needs to break identation.
                                ft.Text(value="Deliverables", expand=True, size=36, weight=ft.FontWeight.W_800),
                                ft.Text(value="Communication Protocols", expand=True, size=24, weight=ft.FontWeight.W_600),
                                ft.Row(
                                    controls=self.checkboxes[:checkboxes_split_point],
                                    alignment=ft.MainAxisAlignment.START,
                                ),
                                ft.Row(
                                    controls=self.checkboxes[checkboxes_split_point:],
                                    alignment=ft.MainAxisAlignment.START,
                                ),
                                ft.Text(value="Device Research", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.device_text_input,
                                ft.Text(value="Vendor Research", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.vendor_text_input,
                                ft.Text(value="Security Research", expand=True, size=24, weight=ft.FontWeight.W_600),
                                self.security_text_input,
                            ]
                        )
                    ]
                )
            )